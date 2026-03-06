"""
╔══════════════════════════════════════════════════════════════════╗
║  rag_engine.py  —  DataMind Platform                            ║
║  THE SMART LIBRARY                                               ║
║                                                                  ║
║  RAG = Retrieval-Augmented Generation.                          ║
║                                                                  ║
║  Imagine a librarian who has read every book in the library.    ║
║  You ask: "What page talks about black holes?"                  ║
║  The librarian doesn't guess from memory — they GO to the shelf,║
║  pull the right chapter, read it, and explain it to you.        ║
║                                                                  ║
║  Step 1 (Ingestion): Read all books → cut into pages → file.   ║
║  Step 2 (Retrieval): Match question to best pages.             ║
║  Step 3 (Generation): AI reads those pages → writes answer.    ║
╚══════════════════════════════════════════════════════════════════╝
"""

# ── IMPORTS ──────────────────────────────────────────────────────────────────

import os           # Operating system tools: file paths, directory listing
import re           # Regular expressions: powerful text pattern matching
import hashlib      # Creates "fingerprints" of text (used for embeddings)
import tempfile     # Creates temporary files/directories (auto-deleted when done)
from typing import List, Dict, Any, Optional, Tuple
import warnings
warnings.filterwarnings("ignore")


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1: CHUNKING UTILITIES
#  "Chunking" = cutting a long document into smaller, searchable pieces.
# ══════════════════════════════════════════════════════════════════════════════

def chunk_text(
    text: str,            # The full document text as one big string
    chunk_size: int = 500,  # Target size of each chunk (in characters)
    overlap: int = 100,   # How many characters to share between adjacent chunks
    separator: str = "\n" # Try to split at newlines first (paragraph boundaries)
) -> List[str]:
    """
    Splits a long document into overlapping chunks.

    WHY OVERLAPPING?
    If we cut perfectly at 500 chars, a sentence like:
    "The temperature rose to 4...
    ...7 degrees overnight."
    ...gets split across two chunks. The AI would miss the full value "47".

    Overlapping means the last 100 characters of chunk 1 = the first 100 of chunk 2.
    No important sentence straddles the boundary undetected.

    Like cutting a newspaper column where you deliberately let the
    scissors overlap slightly, so you never accidentally cut a word in half.
    """
    parts = text.split(separator)
    # Split the whole document by newline characters.
    # A paragraph ends with \n, so this splits by paragraph.
    # e.g. ["Introduction text", "Chapter 1 text", "Chapter 2 text", ...]

    chunks = []   # Will hold our final list of chunks
    current = ""  # The "chunk in progress" that we're currently building

    for part in parts:
        # Loop through each paragraph one by one.

        if len(current) + len(part) <= chunk_size:
            current += part + separator
            # This paragraph FITS in the current chunk. Add it in.
            # Also add back the newline we removed when splitting.

        else:
            # This paragraph would make the chunk TOO LONG. Save the current chunk.
            if current:
                chunks.append(current.strip())
                # .strip() removes leading/trailing whitespace from the saved chunk.

            # Handle paragraphs that are THEMSELVES longer than chunk_size:
            while len(part) > chunk_size:
                chunks.append(part[:chunk_size])
                # Take the first chunk_size characters and save as a chunk.
                # e.g. first 500 characters of a very long paragraph.

                part = part[chunk_size - overlap:]
                # Move forward, but STEP BACK by "overlap" characters.
                # e.g. if chunk_size=500 and overlap=100:
                #   next chunk starts at character 400, not 500.
                # The 100 characters (400–500) appear in BOTH this chunk and the next.

            current = part + separator   # Start a new chunk with the remainder

    if current.strip():
        chunks.append(current.strip())
        # Don't forget the LAST piece after the loop ends.

    return [c for c in chunks if len(c) > 20]
    # Filter out tiny fragments (less than 20 characters are probably just whitespace/newlines).
    # List comprehension: keep only chunks where len(c) > 20.


def extract_text_from_file(file_path: str, file_name: str) -> str:
    """
    Reads text content from different file formats.
    Returns the raw text as one big string.
    """
    ext = file_name.lower().split(".")[-1]
    # Split the filename by "." and take the LAST part = the extension.
    # "report.final.pdf" → ["report", "final", "pdf"] → "pdf"
    # .lower() ensures "PDF" and "pdf" are treated the same.

    if ext in ["txt", "md", "py", "js", "json", "yaml", "yml", "csv"]:
        # These are all plain text formats — just open and read.
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
            # "with open(...)" : safely opens the file and auto-closes it when done.
            # encoding="utf-8" : specify the text encoding (handles international characters)
            # errors="replace" : if a character can't be decoded, use ? instead of crashing.
            # f.read() : reads the ENTIRE file as one big string.

    elif ext == "pdf":
        try:
            import pypdf
            # Import pypdf only when needed (not at the top).
            # This avoids an error if pypdf isn't installed and user doesn't use PDFs.

            reader = pypdf.PdfReader(file_path)
            # PdfReader loads the PDF file.
            # reader.pages is a list: [Page1, Page2, Page3, ...]

            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
            # Loop through every page, extract its text.
            # "or ''" : if a page has no text (scanned image), use empty string.
            # "\n\n".join(...) : put two newlines between each page's text.

        except ImportError:
            return "[PDF parsing requires pypdf. Install with: pip install pypdf]"
            # If pypdf is not installed, return a helpful message instead of crashing.

    elif ext in ["docx", "doc"]:
        try:
            from docx import Document
            # python-docx library for Word documents.

            doc = Document(file_path)
            # Load the Word document.

            return "\n".join(p.text for p in doc.paragraphs)
            # doc.paragraphs = list of Paragraph objects.
            # p.text = the text string inside each paragraph.
            # Join all paragraphs with newlines between them.

        except ImportError:
            return "[DOCX parsing requires python-docx]"

    else:
        # Unknown file type — try to read it as plain text anyway.
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception:
            return f"[Could not parse file type: {ext}]"


def simple_embed(text: str, dim: int = 128) -> List[float]:
    """
    Converts text into a list of 128 numbers called a "vector" or "embedding".

    The idea: similar texts get similar vectors (similar number lists).
    This lets us search by MEANING, not just by keyword.

    How it works:
    1. Extract every small sequence of characters (n-grams) from the text.
    2. "Hash" each n-gram to get a number.
    3. Use that number to fill slots in a 128-number array.
    4. Normalise so all vectors have the same "length" (so comparison is fair).

    NOTE: In production RAG systems, you'd use a neural network embedding
    model like sentence-transformers. This simple version uses character
    hashing — it's less powerful but has zero external dependencies.
    """
    text = text.lower()[:2000]
    # Lowercase all text (so "Apple" and "apple" become the same).
    # [:2000] caps at 2000 characters for speed (very long texts slow things down).

    vec = [0.0] * dim
    # Start with a list of 128 zeros.
    # Like an empty tally sheet with 128 rows.

    for n in range(1, 5):
        # Loop n = 1, 2, 3, 4.
        # n is the "n" in "n-gram" — the length of each character sequence.

        for i in range(len(text) - n + 1):
            # Loop through every possible starting position for an n-gram.
            # For n=2 and text "hello": positions 0,1,2,3 → "he","el","ll","lo"

            gram = text[i:i+n]
            # Cut out the n-gram: n characters starting at position i.
            # e.g. text[1:3] = "el" (characters at index 1 and 2)

            h = int(hashlib.md5(gram.encode()).hexdigest(), 16)
            # gram.encode()  → convert the string to bytes
            # hashlib.md5()  → run the MD5 hash algorithm (creates a "fingerprint")
            # .hexdigest()   → get the result as a hex string like "5d41402abc4b2a76b9719d91"
            # int(..., 16)   → convert from hex (base-16) to a regular integer
            # Now every unique gram has a unique large integer.

            idx = h % dim
            # % (modulo) = remainder after division.
            # e.g. 1234567 % 128 = some number 0 to 127.
            # This maps the large hash number to a slot in our 128-slot vector.

            vec[idx] += 1.0 / n
            # Add a small number to that slot. Shorter n-grams (n=1) contribute more (1/1=1.0).
            # Longer n-grams (n=4) contribute less (1/4=0.25).
            # This weighting makes single characters less dominant than short words.

    mag = (sum(v * v for v in vec) ** 0.5) or 1.0
    # Calculate the "magnitude" (length) of the vector.
    # Pythagoras in N dimensions: sqrt(v1² + v2² + v3² + ... + v128²)
    # "or 1.0" → if mag is 0 (all zeros), use 1.0 to avoid division-by-zero.

    return [v / mag for v in vec]
    # Divide every element by the magnitude.
    # Now the vector has length exactly 1.0 ("unit vector" or "normalised vector").
    # This makes cosine similarity comparisons fair — length doesn't affect the angle.


def cosine_similarity(a: List[float], b: List[float]) -> float:
    """
    Measures how similar two vectors are by the ANGLE between them.
    Returns a value from -1 (completely opposite) to +1 (identical direction).

    WHY ANGLE AND NOT DISTANCE?
    Two embeddings of the same sentence in different lengths would have different
    distances. But their ANGLE would be nearly zero — they point the same direction.
    Cosine similarity is not fooled by the "size" of the text.

    FORMULA: cos(θ) = (a · b) / (|a| × |b|)
    """
    dot = sum(x * y for x, y in zip(a, b))
    # Dot product: multiply corresponding elements and sum them up.
    # zip(a, b) pairs up elements: [(a[0],b[0]), (a[1],b[1]), ...]
    # sum(x * y ...) = a[0]*b[0] + a[1]*b[1] + ... + a[127]*b[127]

    mag_a = sum(x * x for x in a) ** 0.5
    # Magnitude (length) of vector a = sqrt(a[0]² + a[1]² + ...)

    mag_b = sum(y * y for y in b) ** 0.5
    # Magnitude of vector b

    return dot / ((mag_a * mag_b) or 1.0)
    # Divide dot product by the product of magnitudes.
    # "or 1.0" → avoid division by zero if either vector is all zeros.


# ══════════════════════════════════════════════════════════════════════════════
#  CLASS: SimpleVectorStore
#  An in-memory "database" of text chunks with their vector embeddings.
#  Used to store ALL document chunks and search them by similarity.
# ══════════════════════════════════════════════════════════════════════════════

class SimpleVectorStore:
    """
    An in-memory vector store — a searchable list of
    (text, embedding, metadata) tuples.

    When you query it, it computes cosine similarity between the
    query embedding and every stored embedding, then returns the
    top-k most similar results.

    In production this would be ChromaDB, FAISS, or Pinecone.
    This zero-dependency version works perfectly for demos.
    """

    def __init__(self):
        self.documents: List[str] = []
        # List of text strings — one per chunk.
        # e.g. ["Chapter 1: In 1969...", "Chapter 2: The mission...", ...]

        self.embeddings: List[List[float]] = []
        # List of 128-number vectors — one per chunk.
        # Parallel to self.documents: embeddings[i] is the vector for documents[i].

        self.metadatas: List[Dict] = []
        # List of dictionaries — one per chunk.
        # Stores info like which file it came from: {"source": "report.pdf", "chunk_index": 3}

    def add(
        self,
        texts: List[str],          # The text chunks to store
        metadatas: Optional[List[Dict]] = None,   # Optional metadata per chunk
        embed_fn=simple_embed       # The embedding function to use (default = our simple one)
    ):
        """Embeds and stores a batch of text chunks."""
        for i, text in enumerate(texts):
            # enumerate() gives both index (i) and value (text).

            self.documents.append(text)
            # Add the raw text to our documents list.

            self.embeddings.append(embed_fn(text))
            # Compute the embedding vector for this text and store it.
            # embed_fn is a parameter — we can swap in any embedding function.

            meta = metadatas[i] if metadatas else {}
            # If metadata was provided, use the i-th item.
            # Otherwise use an empty dict.

            self.metadatas.append(meta)
            # Store the metadata.

    def query(
        self,
        query_text: str,      # The user's question as a string
        top_k: int = 5,       # How many results to return
        embed_fn=simple_embed  # Embedding function (same one used when adding chunks)
    ) -> List[Dict[str, Any]]:
        """
        Returns the top_k most semantically similar chunks to the query.
        Each result: {"text": "...", "score": 0.87, "metadata": {...}}
        """
        if not self.documents:
            return []
            # Nothing stored yet → return empty list.

        query_emb = embed_fn(query_text)
        # Convert the user's question into a vector using the same embedding function.
        # IMPORTANT: must use the SAME function as when documents were embedded.
        # Comparing embeddings from different functions = garbage results.

        similarities = [
            cosine_similarity(query_emb, doc_emb)
            for doc_emb in self.embeddings
        ]
        # Compute cosine similarity between the query vector and EVERY stored vector.
        # Result: a list of scores, one per stored document.
        # e.g. [0.12, 0.87, 0.34, 0.91, ...] — 0.91 means very similar

        top_indices = sorted(
            range(len(similarities)),   # [0, 1, 2, 3, 4, ...] — all indices
            key=lambda i: similarities[i],   # Sort by the similarity score at that index
            reverse=True                # Highest similarity first
        )[:top_k]
        # sorted() returns indices ordered by their similarity score.
        # [:top_k] keeps only the top 5 indices.
        # e.g. [3, 1, 7, 2, 5] → the 3rd document is most similar.

        return [
            {
                "text": self.documents[i],      # The actual text of this chunk
                "score": round(similarities[i], 4),  # How similar (0–1)
                "metadata": self.metadatas[i],   # Which file, which chunk number
            }
            for i in top_indices
        ]
        # List comprehension: build a list of result dicts for the top matching indices.

    def count(self) -> int:
        return len(self.documents)
        # Simply return how many chunks are stored.

    def clear(self):
        self.documents.clear()
        self.embeddings.clear()
        self.metadatas.clear()
        # Remove ALL stored data from all three lists.
        # Used when the user wants to start fresh with new documents.

    def get_sources(self) -> List[str]:
        """Returns a deduplicated list of all ingested document source filenames."""
        return list({m.get("source", "unknown") for m in self.metadatas})
        # {} creates a SET — sets automatically remove duplicates.
        # m.get("source", "unknown") → get the "source" key, default to "unknown".
        # list() converts the set back to a list.
        # e.g. if 50 chunks all came from "report.pdf", this returns just ["report.pdf"].


# ══════════════════════════════════════════════════════════════════════════════
#  CLASS: RAGEngine
#  The full pipeline: ingest documents → retrieve relevant chunks → build prompts
# ══════════════════════════════════════════════════════════════════════════════

class RAGEngine:
    """
    The full Retrieval-Augmented Generation pipeline.

    Two phases:
    PHASE 1 — Ingestion:
        Upload file → extract text → chunk → embed → store in vector store

    PHASE 2 — Querying:
        Ask question → embed question → find similar chunks → build prompt → stream answer
    """

    def __init__(self):
        self.store = SimpleVectorStore()
        # The vector store where all document chunks (and their embeddings) live.

        self.ingested_files: List[Dict] = []
        # A log of which files have been ingested.
        # e.g. [{"file_name": "report.pdf", "chunks": 42, "status": "success"}]

    # ── Document Ingestion ────────────────────────────────────────────────────

    def ingest_file(self, file_path: str, file_name: str, chunk_size: int = 600, overlap: int = 100) -> Dict:
        """
        Full ingestion pipeline for one file.

        CSV files get SPECIAL treatment:
          Each row becomes its own searchable chunk.
          e.g. "Title: Inception | Year: 2010 | Rating: 8.8"
          This lets RAG answer questions like "top rated movies" by finding
          individual rows instead of hunting through one big blob of text.

        All other files (PDF, TXT, DOCX etc.) use the standard
        overlapping-chunk approach.
        """
        ext = file_name.lower().split(".")[-1]

        # ── Special path: CSV row-by-row ingestion ────────────────────────
        if ext == "csv":
            try:
                import pandas as pd
                df = pd.read_csv(file_path)

                chunks = []
                for _, row in df.iterrows():
                    # Convert every row to a "col: val | col: val | ..." string.
                    # Skips empty/NaN values so chunks stay clean.
                    row_text = " | ".join(
                        f"{col}: {val}"
                        for col, val in row.items()
                        if str(val).strip() not in ("", "nan", "None", "NaN")
                    )
                    if row_text.strip():
                        chunks.append(row_text)

                if not chunks:
                    return {"status": "error", "message": "No rows extracted from CSV", "chunks": 0}

                metadatas = [
                    {"source": file_name, "chunk_index": i,
                     "total_chunks": len(chunks), "type": "csv_row"}
                    for i in range(len(chunks))
                ]
                self.store.add(chunks, metadatas=metadatas)

                # Add a header summary so the AI knows all column names.
                header_chunk = (
                    f"Dataset '{file_name}' columns: {', '.join(df.columns.tolist())}. "
                    f"Total rows: {len(df)}."
                )
                self.store.add(
                    [header_chunk],
                    metadatas=[{"source": file_name, "chunk_index": -1, "type": "csv_header"}]
                )

                info = {
                    "file_name": file_name,
                    "total_chars": sum(len(c) for c in chunks),
                    "chunks": len(chunks) + 1,
                    "status": "success",
                }
                self.ingested_files.append(info)
                return info

            except Exception:
                pass  # Fall through to regular text ingestion if pandas fails

        # ── Standard path: text extraction + overlapping chunks ───────────
        raw_text = extract_text_from_file(file_path, file_name)

        if not raw_text.strip():
            return {"status": "error", "message": "No text extracted", "chunks": 0}

        chunks = chunk_text(raw_text, chunk_size=chunk_size, overlap=overlap)

        metadatas = [
            {"source": file_name, "chunk_index": i, "total_chunks": len(chunks)}
            for i in range(len(chunks))
        ]
        self.store.add(chunks, metadatas=metadatas)

        info = {
            "file_name": file_name,
            "total_chars": len(raw_text),
            "chunks": len(chunks),
            "status": "success",
        }
        self.ingested_files.append(info)
        return info

    def ingest_text(self, text: str, source_name: str = "manual", chunk_size: int = 600) -> Dict:
        """
        Ingest raw text directly (not from a file).
        Useful for adding DataFrame summaries or user-pasted text to the knowledge base.
        """
        chunks = chunk_text(text, chunk_size=chunk_size)
        # Chunk the text (same as for files, just without the file reading step).

        metadatas = [{"source": source_name, "chunk_index": i} for i in range(len(chunks))]
        # Build metadata — here source_name is the label we give this text.

        self.store.add(chunks, metadatas=metadatas)
        return {"source": source_name, "chunks": len(chunks), "status": "success"}

    # ── Retrieval ─────────────────────────────────────────────────────────────

    def retrieve(self, query: str, top_k: int = 5) -> List[Dict]:
        """
        Fetches the most relevant context chunks for the user's question.
        Delegates to the vector store's query() method.
        """
        return self.store.query(query, top_k=top_k)
        # Simply calls the vector store's query function.
        # The vector store handles embedding the question and computing similarities.

    def build_prompt(self, query: str, top_k: int = 5) -> Tuple[str, List[Dict]]:
        """
        The KEY function: builds the "augmented" prompt for the AI.

        1. Retrieve the most relevant chunks.
        2. Format them with source labels.
        3. Prepend a system instruction: "Answer ONLY from context."
        4. Append the user's question.

        This structured prompt is what prevents hallucination:
        the AI has REAL evidence to work from.

        Returns: (prompt_string, list_of_retrieved_chunks)
        """
        retrieved = self.retrieve(query, top_k=top_k)
        # Get the top relevant chunks from the vector store.

        if not retrieved:
            context_block = "[No relevant documents found in the knowledge base.]"
            # If nothing was found (empty vector store, or unrelated question),
            # use a placeholder message so the AI knows there's no context.
        else:
            context_parts = []
            for i, chunk in enumerate(retrieved, 1):
                # enumerate(retrieved, 1) starts counting at 1 (not 0), for readable labels.

                source = chunk["metadata"].get("source", "unknown")
                # Which file did this chunk come from?

                score = chunk["score"]
                # How similar was this chunk to the question? (0.0 to 1.0)

                context_parts.append(
                    f"[Context {i} | Source: {source} | Relevance: {score:.3f}]\n{chunk['text']}"
                )
                # Format each chunk as a labelled block.
                # e.g. "[Context 1 | Source: report.pdf | Relevance: 0.872]
                #       Chapter 3 describes the melting rate of arctic ice..."

            context_block = "\n\n---\n\n".join(context_parts)
            # Join all context blocks with separators between them.
            # "---" acts like a horizontal rule — clearly separates one chunk from the next.

        prompt = f"""You are DataMind's knowledge base assistant. Answer the question using ONLY the provided context.
If the answer is not in the context, say "I don't have enough context to answer this."
Always cite which context block(s) your answer is based on.

CONTEXT:
{context_block}

QUESTION: {query}

ANSWER:"""
        # f-string with triple quotes: allows multi-line strings.
        # {context_block} gets replaced with the actual retrieved text.
        # {query} gets replaced with the user's actual question.
        # The "ONLY" instruction is critical — it forces grounded answers.
        # "Always cite" makes the AI say "According to Context 1..." in its answer.

        return prompt, retrieved
        # Return BOTH the built prompt AND the raw retrieved chunks.
        # The app uses retrieved to show "here's what was found" to the user.

    # ── Status & Maintenance ──────────────────────────────────────────────────

    def get_status(self) -> Dict:
        """Returns a summary of what's stored in the knowledge base."""
        return {
            "total_chunks":    self.store.count(),         # Total chunks stored
            "files_ingested":  len(self.ingested_files),   # Number of files processed
            "sources":         self.store.get_sources(),   # Unique filenames
            "files":           self.ingested_files,        # Full log of files
        }

    def clear(self):
        """Wipe everything — start fresh with an empty knowledge base."""
        self.store.clear()          # Remove all vectors and text from the store
        self.ingested_files.clear() # Clear the ingestion log
        # .clear() is a list method that empties the list in-place (faster than = []).
